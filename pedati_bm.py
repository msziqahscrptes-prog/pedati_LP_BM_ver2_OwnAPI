import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Rancangan Mengajar Utama PEDATI", layout="wide")
st.title("🎓 Penjana Rancangan Mengajar PEDATI")

# --- INPUT KUNCI API (DI ATAS TOPIK) ---
user_api_key = st.text_input(
    "🔑 Masukkan Kunci Rahsia API Gemini Anda:", 
    type="password", 
    help="Dapatkan kunci API anda dari Google AI Studio menggunakan akaun Gmail anda."
)

# Fungsi pembantu untuk memeriksa dan memuatkan model secara dinamik
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except Exception as e:
        st.error(f"Ralat Kunci API atau Masalah Sambungan: {str(e)}")
        return None
    return "models/gemini-1.5-flash"  # Sandaran Lalai


# Memproses pengesahan model
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"Sistem disambungkan. Model Aktif: {selected_model_name}")
else:
    st.warning("⚠️ Sila masukkan Kunci API Gemini peribadi anda di atas untuk bermula.")


def generate_pedati_plan(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    # Arahan Prompt dalam Bahasa Melayu sepenuhnya
    prompt = f"""
    Topik: {topic}. Kod Silibus: {syllabus}. Konteks Tambahan: {extra_context}.
    Janakan rancangan mengajar dalam Bahasa Melayu sepenuhnya. 
    Jangan guna perkataan 'Murid', gantikan ia dengan perkataan 'Pelajar'.
    Gunakan nama peringkat PEDATI yang tepat seperti berikut:
    P [Pengetahuan Sedia Ada], E [Empati / Penglibatan], D [Daya Upaya / Perkembangan], A [Aplikasi], T [Taksir], I [Improvisasi / Penambahbaikan].

    Strukturkan output dengan penanda (markers) berikut untuk pembetukan jadual/kotak:
    SECTION: OBJEKTIF PELAJARAN
    [4 mata/poin]
    SECTION: HASIL PELAJARAN
    [4 mata/poin]
    SECTION: KRITERIA KEJAYAAN
    [4 mata/poin]
    SECTION: PRASYARAT
    [1 mata/poin]
    SECTION: KATA KUNCI
    [6 item]
    SECTION: KBAT
    [mana-mana 4 domain utama dalam Taksonomi Bloom]
    SECTION: KEWARGANEGARAAN DIGITAL
    [4 mata/poin mengenai penggunaan sumber dalam talian seperti saluran youtube, aplikasi canva, chromebook, atau peranti digital]

    SECTION: PERINGKAT PEDATI
    STAGE: P [Pengetahuan Sedia Ada] | CB: [Aktiviti Guru/Fasilitator] | SB: [Aktiviti Murid]
    STAGE: E [Empati / Penglibatan] | CB: [Aktiviti Guru/Fasilitator] | SB: [Aktiviti Murid]
    STAGE: D [Daya Upaya / Perkembangan] | CB: [Aktiviti Guru/Fasilitator] | SB: [Aktiviti Murid]
    STAGE: A [Aplikasi] | CB: [Aktiviti Guru/Fasilitator] | SB: [Aktiviti Murid]
    STAGE: T [Taksir] | CB: [Aktiviti Guru/Fasilitator] | SB: [Aktiviti Murid]
    STAGE: I [Improvisasi / Penambahbaikan] | CB: [Aktiviti Guru/Fasilitator] | SB: [Aktiviti Murid]
    """
    try:
        response = model.generate_content(prompt)
        
        # Pemeriksaan keselamatan kandungan respons
        if response.candidates and response.candidates[0].content.parts:
            return response.text
        else:
            return "⚠️ AI mengembalikan respons kosong. Anda mungkin menekan butang terlalu cepat sehingga melebihi had kuota seminit (15 permintaan/minit). Sila tunggu 60 saat dan cuba lagi."
            
    except Exception as e:
        return f"Ralat Sistem: {str(e)}"


def create_word_export(topic, syllabus, text):
    doc = Document()
    doc.add_heading(f'Rancangan Mengajar: {topic} ({syllabus})', 0)

    # 1. Jadual Pentadbiran Atas (Admin Header Table)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Minggu No :", "Tarikh:"], ["Bil. Pelajar:", "Hari:"], ["Tempat / No. Makmal:", "Durasi (minit):"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    # 2. Jadual Sumber & Bahan
    doc.add_heading("Sumber & Bahan Bantu Mengajar", level=1)
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    res_table.cell(0, 0).text = "Papan pintar (Smart board), Chromebook, Meja tulis, Projektor, Perkongsian skrin dengan komputer riba"

    # 3. Penguraian Kandungan & Pembentukan Jadual
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip()
        content_lines = lines[1:]
        doc.add_heading(title.title(), level=1)

        if "|" in section and "PEDATI" in title.upper():
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Peringkat (PEDATI)', 'Pensyarah(Tutor)', 'Pelajar(Penuntut)'

            for line in content_lines:
                if "|" in line:
                    p = line.split("|")
                    row = table.add_row().cells
                    row[0].text = p[0].split(":")[-1].strip()
                    row[1].text = p[1].split(":")[-1].strip()
                    row[2].text = p[2].split(":")[-1].strip()
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "\n".join([l.strip() for l in content_lines if l.strip()])

    # 4. Halaman Kelulusan HKP (HOD Approval)
    doc.add_page_break()
    doc.add_heading("Kelulusan & Ulasan Pengetua / Ketua Jabatan", level=1)
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Ulasan / Catatan"
    hod_table.cell(0, 1).text = "Tandatangan / Cop Rasmi"
    hod_table.rows[1].height = Pt(30)
    hod_table.cell(2, 0).text = "Tarikh:"
    hod_table.cell(2, 1).text = "Nama:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 2. GUI UTAMA ---
st.write("---") # Garis pemisah bawah bar API
c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Topik Pelajaran:")
with c2: u_syllabus = st.text_input("Kod Silibus:")
u_extra = st.text_area("Konteks Khas / Kata Kunci Tambahan (Opsional):")

if st.button("🚀 JANAKAN RANCANGAN MENGAJAR PEDATI"):
    if not user_api_key:
        st.error("❌ Sila masukkan Kunci API Gemini anda di bahagian atas sebelum menjana.")
    elif not u_topic or not u_syllabus:
        st.error("❌ Sila isi maklumat Topik Pelajaran dan Kod Silibus.")
    else:
        with st.spinner("AI sedang membina rancangan mengajar PEDATI anda..."):
            result = generate_pedati_plan(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['pedati_out_ms'] = result

if 'pedati_out_ms' in st.session_state:
    st.divider()
    st.text_area("Pratonton AI", st.session_state['pedati_out_ms'], height=300)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['pedati_out_ms'])

    st.download_button("📥 Muat Turun Fail Word (.docx)", doc_file, f"PEDATI_{u_topic}.docx")

# --- SEKSYEN KAKI (FOOTER) ---
st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>Smart PEDATI Lesson Plan AI-Generator (Versi BM) v1.0</b></p>
        <p>Developed & Conceptualized by: <b> Hajah Nurul Haziqah @ Hjh Hartini Hj Nordin </b></p>
        <p>© 2026 BSc.M.(H) Computer Science, Strathclyde University</p>
    </div>
    """,
    unsafe_allow_html=True
)
